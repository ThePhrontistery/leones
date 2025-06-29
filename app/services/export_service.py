"""
Servicio para exportar documentos Markdown a Word (.docx) y PDF con tabla de contenido automática.
"""
from typing import Any
from io import BytesIO
from docx import Document
from docx.shared import Pt
import markdown2
from bs4 import BeautifulSoup
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fpdf import FPDF
import re

async def markdown_to_docx(markdown_text: str) -> BytesIO:
    """
    Convierte un texto Markdown a un archivo Word (.docx) en memoria.
    Incluye una tabla de contenido automática y aplica estilos a títulos, listas y negritas.
    """
    # Convertir Markdown a HTML
    html = markdown2.markdown(markdown_text)
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()

    # Añadir título para la tabla de contenido
    doc.add_heading("Tabla de contenido", level=1)
    # Instrucción para el usuario (en un párrafo aparte, antes del campo TOC)
    doc.add_paragraph("(Pulsa Ctrl + A -o Ctrl + E en algunas configuraciones regionales- y luego F9 para actualizar la tabla de contenido y ver los títulos)", style="Intense Quote")

    # Insertar tabla de contenido automática (campo TOC) en un párrafo propio, sin texto adicional
    p = doc.add_paragraph()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    r = p.add_run()._r
    r.append(fldChar_begin)
    r.append(instrText)
    r.append(fldChar_separate)
    r.append(fldChar_end)
    p.alignment = 0  # left

    # Salto de página tras la TOC
    doc.add_page_break()

    def clean_heading(text: str) -> str:
        # Elimina patrones tipo {#...} al final del texto
        import re
        return re.sub(r'\s*\{#.*?\}\s*$', '', text).strip()

    def add_element(element: Any):
        if element.name == "h1":
            doc.add_heading(clean_heading(element.get_text()), level=1)
        elif element.name == "h2":
            doc.add_heading(clean_heading(element.get_text()), level=2)
        elif element.name == "h3":
            doc.add_heading(clean_heading(element.get_text()), level=3)
        elif element.name == "ul":
            for li in element.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(), style="List Bullet")
        elif element.name == "ol":
            for li in element.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(), style="List Number")
        elif element.name == "p":
            p = doc.add_paragraph()
            for child in element.children:
                if getattr(child, "name", None) == "strong":
                    run = p.add_run(clean_heading(child.get_text()))
                    run.bold = True
                else:
                    p.add_run(str(child))
        elif element.name == "strong":
            run = doc.add_paragraph().add_run(clean_heading(element.get_text()))
            run.bold = True
        elif element.name is None:
            doc.add_paragraph(element)
        # Puedes agregar más reglas para otros elementos

    for el in soup.body.contents if soup.body else soup.contents:
        add_element(el)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def parse_markdown_lines(markdown_text: str):
    """
    Parsea el markdown línea a línea y detecta títulos, listas y negritas.
    Devuelve una lista de tuplas (tipo, texto).
    """
    lines = markdown_text.splitlines()
    parsed = []
    for line in lines:
        line = line.rstrip()
        if not line.strip():
            parsed.append(("blank", ""))
        elif re.match(r"^# ", line):
            parsed.append(("h1", re.sub(r"^# ", "", line)))
        elif re.match(r"^## ", line):
            parsed.append(("h2", re.sub(r"^## ", "", line)))
        elif re.match(r"^### ", line):
            parsed.append(("h3", re.sub(r"^### ", "", line)))
        elif re.match(r"^[-*+] ", line):
            parsed.append(("ul", re.sub(r"^[-*+] ", "", line)))
        elif re.match(r"^\d+\. ", line):
            parsed.append(("ol", re.sub(r"^\d+\. ", "", line)))
        else:
            parsed.append(("p", line))
    return parsed

def asciify(text: str) -> str:
    """
    Reemplaza caracteres Unicode problemáticos y acentos/ñ por equivalentes ASCII seguros para FPDF.
    """
    replacements = {
        "á": "a", "é": "e", "í": "i", "ó": "o", "ú": "u",
        "Á": "A", "É": "E", "Í": "I", "Ó": "O", "Ú": "U",
        "ñ": "n", "Ñ": "N",
        "ü": "u", "Ü": "U",
        "“": '"', "”": '"', "‘": "'", "’": "'",
        "–": "-", "—": "-", "…": "...", "•": "-", "·": "-",
        "→": ">", "←": "<", "«": '"', "»": '"',
        "©": "(c)", "®": "(R)", "™": "(TM)", "°": " deg ",
        "€": "EUR", "£": "GBP", "¥": "YEN", "§": "S", "¶": "P",
        # Agrega más si es necesario
    }
    return ''.join(replacements.get(c, c if ord(c) < 128 else '?') for c in text)

def clean_ascii(text: str) -> str:
    """
    Elimina cualquier carácter no ASCII (fuera de 32-126) y saltos de línea.
    """
    return ''.join(c for c in text if 32 <= ord(c) <= 126)

def safe_multicell(pdf, h, text):
    """
    Llama a pdf.multi_cell(0, h, text) solo si el texto es seguro.
    Si el texto es vacío, un solo carácter, o contiene caracteres no imprimibles, imprime '- item seguro'.
    """
    safe = text.strip()
    # Elimina caracteres de control
    safe = ''.join(c for c in safe if 32 <= ord(c) <= 126)
    if len(safe) < 2:
        pdf.multi_cell(0, h, "- item seguro")
    else:
        pdf.multi_cell(0, h, safe)

async def markdown_to_pdf(markdown_text: str) -> bytes:
    """
    Convierte Markdown a PDF usando FPDF2 (ASCII, sin desbordes, soporte negrita básica).
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    parsed = parse_markdown_lines(markdown_text)
    for tipo, texto in parsed:
        texto = re.sub(r'\s*\{#.*?\}\s*$', '', texto).strip()  # Limpia IDs
        texto = asciify(texto)
        texto = clean_ascii(texto)
        if tipo == "h1":
            pdf.set_font("Arial", "B", 18)
            safe_multicell(pdf, 10, texto)
            pdf.set_font("Arial", size=12)
        elif tipo == "h2":
            pdf.set_font("Arial", "B", 15)
            safe_multicell(pdf, 8, texto)
            pdf.set_font("Arial", size=12)
        elif tipo == "h3":
            pdf.set_font("Arial", "B", 13)
            safe_multicell(pdf, 7, texto)
            pdf.set_font("Arial", size=12)
        elif tipo == "ul" or tipo == "ol":
            pdf.cell(8)
            for subline in texto.split("\n"):
                clean_text = subline.strip()
                clean_text = clean_ascii(clean_text)
                if not clean_text or len(clean_text.replace('-', '').replace(' ', '')) < 2:
                    try:
                        pdf.cell(0, 7, "- item seguro de lista", ln=1)
                    except Exception:
                        pdf.multi_cell(0, 7, "- item seguro de lista")
                    continue
                # Divide en líneas de máximo 90 caracteres para evitar errores de ancho
                maxlen = 90
                lines = [clean_text[i:i+maxlen] for i in range(0, len(clean_text), maxlen)]
                for idx, line in enumerate(lines):
                    try:
                        if len(line) < maxlen:
                            pdf.cell(0, 7, f"- {line}", ln=1)
                        else:
                            pdf.multi_cell(0, 7, f"- {line}")
                    except Exception:
                        try:
                            pdf.multi_cell(0, 7, f"- {line}")
                        except Exception:
                            pdf.cell(0, 7, "- item seguro de lista", ln=1)
        elif tipo == "p":
            # Soporte negrita básica: **texto**
            bold_parts = re.split(r'(\*\*[^*]+\*\*)', texto)
            for part in bold_parts:
                part = clean_ascii(part)
                if part.startswith('**') and part.endswith('**'):
                    pdf.set_font("Arial", "B", 12)
                    pdf.write(7, part[2:-2])
                    pdf.set_font("Arial", size=12)
                else:
                    pdf.write(7, part)
            pdf.ln(7)
        elif tipo == "blank":
            pdf.ln(3)
    output_bytes = pdf.output(dest="S")
    return bytes(output_bytes)
